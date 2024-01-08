from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt
import pandas as pd


def singleStep(docu, name, awardName, serialNum):
    para1 = docu.add_paragraph()
    r1 = para1.add_run("\n " + name + " ")
    r1.underline = True
    r1.font.name = u"隶书"
    r1.font.size = Pt(48)
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), u'隶书')

    r2 = para1.add_run("老师：")
    r2.font.name = u"华文中宋"
    r2.font.size = Pt(36)
    r2.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')

    para2 = docu.add_paragraph("荣获2022年度大学生创新创业")
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in para2.runs:
        r.font.name = u"华文中宋"
        r.font.size = Pt(36)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')

    para3 = docu.add_paragraph(awardName)
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for r in para3.runs:
        r.font.name = u"华文新魏"
        r.font.size = Pt(36)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文新魏')

    para4 = docu.add_paragraph("证书编号：" + serialNum)
    for r in para4.runs:
        r.font.name = u"华文中宋"
        r.font.size = Pt(18)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')

    para5 = docu.add_paragraph("　上海理工大学　")
    para5.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for r in para5.runs:
        r.font.name = u"华文中宋"
        r.font.size = Pt(26)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')

    para6 = docu.add_paragraph("二〇二二年十二月")
    para6.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for r in para6.runs:
        r.font.name = u"华文中宋"
        r.font.size = Pt(26)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')

    para6.runs[-1].add_break(WD_BREAK.PAGE)
    for para in docu.paragraphs:
        for run in para.runs:
            run.bold = True
    # docu.save("0.docx")


if __name__ == '__main__':
    docu = Document()
    default_section = docu.sections[0]
    default_section.orientation = WD_ORIENTATION.LANDSCAPE
    default_section.page_width, default_section.page_height = default_section.page_height, default_section.page_width
    data = pd.read_excel("clearedData.xlsx")
    name = data["name"]
    awardName = data["awardName"]
    finalSerialNum = data["finalSerialNum"]
    for i, j, k in zip(name, awardName, finalSerialNum):
        singleStep(docu, i, j, k)
    docu.save("获奖信息.docx")
