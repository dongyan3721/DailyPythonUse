from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import pandas as pd


def createSingleDocu(name, code, position):
    docu = Document()

    para0 = docu.add_paragraph("志愿服务证明").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    para1 = docu.add_paragraph("\u3000兹有")
    para1.add_run(" 基础 ").underline = True
    para1.add_run("学院")
    para1.add_run("   ").underline = True
    para1.add_run(name).underline = True
    para1.add_run("   ").underline = True
    para1.add_run("同学，（学号：")
    para1.add_run("  ").underline = True
    para1.add_run(code).underline = True
    para1.add_run("   ").underline = True
    para1.add_run('）')
    para1.add_run("，于2022年8月基础学院搬迁校区时，担任")

    para1.add_run("   ").underline = True
    para1.add_run(position).underline = True
    para1.add_run("   ").underline = True
    para1.add_run("志愿者一职，期间态度认真，服务积极，表现突出，为完成校区搬迁工作贡献出自己的一份青年力量。")
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para3 = docu.add_paragraph("\u3000特此证明！")
    para3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para2 = docu.add_paragraph("上海理工大学基础学院团委\n")
    para2.add_run("2022年8月29日\n")
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    for para in docu.paragraphs:
        for run in para.runs:
            run.font.name = u"宋体"
            run.font.size = Pt(14.5)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    for run in docu.paragraphs[0].runs:
        run.bold = True
        run.font.name = u"宋体"
        run.font.size = Pt(20)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    docu.paragraphs[0].space_after = Pt(15)

    docu.save(code + name + '.docx')


if __name__ == '__main__':
    # createSingleDocu("XXX", "2135XXXXXX", "核酸检测", "411624001089522364")
    data = pd.read_excel('specializedData.xlsx')
    l1 = list(data['姓名'])
    l2 = list(data['学号'])
    l3 = list(data['职位'])
    # l4 = list(data['志愿者编号'])
    # for n, c, p, v in zip(l1, l2, l3, l4):
    #     createSingleDocu(n, str(c), p, str(v))
    for n, c, p in zip(l1, l2, l3):
        createSingleDocu(n, str(int(c)), p)

